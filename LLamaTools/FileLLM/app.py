import streamlit as st
import os
import subprocess
import torch
import asyncio
import fitz
import pdfplumber
import pytesseract
from pdf2image import convert_from_bytes
import pandas as pd
import json
import xml.etree.ElementTree as ET
import io
from docx import Document
from bs4 import BeautifulSoup
from langchain_ollama import OllamaLLM
from langchain_core.prompts import ChatPromptTemplate
from urllib.parse import urljoin
import requests
from readability import Document as ReadabilityDocument
import re
import psutil
from crawl4ai import AsyncWebCrawler, BrowserConfig, CrawlerRunConfig, CacheMode
import platform

# Check if GPU is available
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

# Check if Ollama is running, if not, start it
def is_ollama_running():
    try:
        output = subprocess.check_output('tasklist', shell=True).decode()
        return 'ollama.exe' in output
    except subprocess.CalledProcessError:
        return False

if not is_ollama_running():
    subprocess.Popen(['start', 'ollama'], shell=True)

# Streamlit UI
st.title('Ollama Llama3 Model')

# Initialize session state
if 'uploaded_file' not in st.session_state:
    st.session_state.uploaded_file = None
if 'website_url' not in st.session_state:
    st.session_state.website_url = None
if 'website_toggle' not in st.session_state:
    st.session_state.website_toggle = False
if 'scraped_results' not in st.session_state:
    st.session_state.scraped_results = []
if 'conversation_history' not in st.session_state:
    st.session_state.conversation_history = []
if 'scraping_started' not in st.session_state:
    st.session_state.scraping_started = False

# Function to toggle website input and file upload
def toggle_website_input():
    if st.session_state.website_toggle:
        st.session_state.uploaded_file = None
    else:
        st.session_state.website_url = None
    st.session_state.scraping_started = False

# Create a toggle for website input
st.checkbox('Enter Website URL', key='website_toggle', on_change=toggle_website_input)

# Show website input or file uploader based on toggle
if st.session_state.website_toggle:
    st.session_state.website_url = st.text_input('Enter Website URL')
else:
    st.session_state.uploaded_file = st.file_uploader(
        'Choose a file to upload',
        type=['txt', 'pdf', 'docx', 'xlsx', 'json', 'xml', 'html', 'csv'],
        key='file_uploader'
    )

uploaded_file = st.session_state.uploaded_file
website_url = st.session_state.website_url

@st.cache_resource
def load_model():
    template = (
        "You are tasked with answering questions based on the following document content: {document_content}. "
        "Please provide accurate and concise answers. Here is the conversation history: {conversation_history}. "
        "Now, answer the following question: {question}"
    )
    model = OllamaLLM(model="llama3", device=device)
    return model, template

# Add a submit button for URL scraping
if website_url:
    if st.button('Submit URL'):
        st.session_state.scraping_started = True
        st.write(f'Scraping website: {website_url}')

# Scraping logic
if st.session_state.scraping_started and website_url:
    async def scrape_and_extract_text(urls, max_concurrent=10):
        browser_config = BrowserConfig(
            headless=True,
            verbose=False,
            extra_args=["--disable-gpu", "--disable-dev-shm-usage", "--no-sandbox"],
        )
        crawl_config = CrawlerRunConfig(cache_mode=CacheMode.BYPASS)

        crawler = AsyncWebCrawler(config=browser_config)
        await crawler.start()

        extracted_texts = []
        process = psutil.Process(os.getpid())
        peak_memory = 0

        def log_memory(prefix: str = ""):
            nonlocal peak_memory
            current_mem = process.memory_info().rss
            if current_mem > peak_memory:
                peak_memory = current_mem
            print(f"{prefix} Current Memory: {current_mem // (1024 * 1024)} MB, Peak: {peak_memory // (1024 * 1024)} MB")

        try:
            for i in range(0, len(urls), max_concurrent):
                batch = urls[i:i + max_concurrent]
                tasks = [crawler.arun(url=url, config=crawl_config, session_id=f"session_{i+j}") for j, url in enumerate(batch)]
                log_memory(prefix=f"Before batch {i//max_concurrent + 1}: ")
                results = await asyncio.gather(*tasks)
                log_memory(prefix=f"After batch {i//max_concurrent + 1}: ")

                for result in results:
                    if result.success:
                        print(f"Successfully crawled: {result.url}")
                        try:
                            doc = ReadabilityDocument(result.html)
                            content = doc.summary()
                            soup = BeautifulSoup(content, 'lxml')
                            text = soup.get_text(separator='\n').strip()
                            text = re.sub(r'http\S+', '', text)
                            extracted_texts.append(text)
                        except Exception as e:
                            print(f"Error extracting text from {result.url}: {e}")
                    else:
                        print(f"Failed: {result.url} - Error: {result.error_message}")
        finally:
            await crawler.close()
            log_memory(prefix="Final: ")

        return extracted_texts

    async def main():
        sitemap = [website_url]
        st.session_state.scraped_results = await scrape_and_extract_text(sitemap)
        st.text_area('Scraped Results', '\n'.join(st.session_state.scraped_results), height=150)

        # Load the model and template
        model, template = load_model()

        # Function to get response from the model
        async def get_response(prompt):
            st.session_state.conversation_history.append({'role': 'user', 'content': prompt})
            prompt_template = ChatPromptTemplate.from_template(template)
            chain = prompt_template | model
            response = await asyncio.to_thread(chain.invoke, {
                "document_content": st.session_state.scraped_results,
                "conversation_history": st.session_state.conversation_history,
                "question": prompt
            })
            st.session_state.conversation_history.append({'role': 'assistant', 'content': response})
            return response.strip()

        user_input = st.text_input('Ask your question:')
        if st.button('Submit Question'):
            response = await get_response(user_input)
            st.write(response)

    if platform.system() == 'Windows':
        asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())
    asyncio.run(main())

# OCR-based text extraction (if PyMuPDF fails)
def extract_text_with_ocr(file_content):
    try:
        images = convert_from_bytes(file_content)
        text = "\n".join([pytesseract.image_to_string(img, lang="eng") for img in images])
        return text if text.strip() else None
    except Exception as e:
        st.error(f"OCR extraction failed: {e}")
        return None

# PdfPlumber fallback for extracting text
def extract_text_with_pdfplumber(file_content):
    try:
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            text = "\n".join([page.extract_text() or "" for page in pdf.pages])
            return text if text.strip() else None
    except Exception as e:
        st.error(f"PdfPlumber extraction failed: {e}")
        return None

# DOCX extraction function
def extract_text_from_docx(file_content):
    try:
        doc = Document(io.BytesIO(file_content))
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        st.error(f"Error reading DOCX file: {e}")
        return None

# TXT extraction function
def extract_text_from_txt(file_content):
    try:
        return file_content.decode("utf-8").strip()
    except UnicodeDecodeError:
        return file_content.decode("latin-1").strip()

# CSV extraction function
def extract_text_from_csv(file_content):
    try:
        df = pd.read_csv(io.BytesIO(file_content))
        return df.to_string(index=False)
    except Exception as e:
        st.error(f"Error reading CSV file: {e}")
        return None

# XLSX extraction function
def extract_text_from_xlsx(file_content):
    try:
        df = pd.read_excel(io.BytesIO(file_content), sheet_name=None)
        content = []
        for sheet, data in df.items():
            content.append(f"Sheet: {sheet}\n{data.to_string(index=False)}\n")
        return "\n".join(content)
    except Exception as e:
        st.error(f"Error reading Excel file: {e}")
        return None

# JSON extraction function
def extract_text_from_json(file_content):
    try:
        json_data = json.loads(file_content.decode("utf-8"))
        return json.dumps(json_data, indent=2)
    except Exception as e:
        st.error(f"Error reading JSON file: {e}")
        return None

# XML extraction function
def extract_text_from_xml(file_content):
    try:
        tree = ET.ElementTree(ET.fromstring(file_content.decode("utf-8")))
        root = tree.getroot()
        return "\n".join([ET.tostring(child, encoding="unicode") for child in root])
    except Exception as e:
        st.error(f"Error reading XML file: {e}")
        return None

# HTML extraction function
def extract_text_from_html(file_content):
    try:
        soup = BeautifulSoup(file_content, "html.parser")
        return soup.get_text()
    except Exception as e:
        st.error(f"Error reading HTML file: {e}")
        return None

# Function to extract text from various file types
def extract_text_from_file(uploaded_file):
    if uploaded_file is None:
        raise ValueError("No file uploaded.")

    file_content = uploaded_file.read()
    if not file_content:
        raise ValueError("Uploaded file is empty. Please upload a valid file.")

    file_type = uploaded_file.type

    if file_type == "application/pdf":
        try:
            pdf_document = fitz.open(stream=file_content, filetype="pdf")
            text = "\n".join([pdf_document.load_page(i).get_text("text") for i in range(len(pdf_document))])
            if text.strip():
                return text
        except Exception as e:
            st.warning(f"PyMuPDF failed: {e}")

        st.warning("No selectable text found. Trying OCR...")
        text = extract_text_with_ocr(file_content)
        if text:
            return text

        st.warning("OCR failed. Trying PdfPlumber as a last resort...")
        return extract_text_with_pdfplumber(file_content)

    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_text_from_docx(file_content)

    elif file_type == "text/plain":
        return extract_text_from_txt(file_content)

    elif file_type == "text/csv":
        return extract_text_from_csv(file_content)

    elif file_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        return extract_text_from_xlsx(file_content)

    elif file_type == "application/json":
        return extract_text_from_json(file_content)

    elif file_type == "text/xml":
        return extract_text_from_xml(file_content)

    elif file_type == "text/html":
        return extract_text_from_html(file_content)

    else:
        raise ValueError("Unsupported file type.")

# Main execution when a file is uploaded
if uploaded_file is not None:
    try:
        file_content = extract_text_from_file(uploaded_file)
        st.success("File uploaded successfully!")

        # Display extracted text
        st.text_area('Extracted Text', file_content, height=300)

        model, template = load_model()
        st.session_state.conversation_history = []

        # Function to get response from the model
        async def get_response(prompt):
            st.session_state.conversation_history.append({'role': 'user', 'content': prompt})
            prompt_template = ChatPromptTemplate.from_template(template)
            chain = prompt_template | model
            response = await asyncio.to_thread(chain.invoke, {
                "document_content": file_content,
                "conversation_history": st.session_state.conversation_history,
                "question": prompt
            })
            st.session_state.conversation_history.append({'role': 'assistant', 'content': response})
            return response.strip()

        user_input = st.text_input('Ask your question:')
        if st.button('Submit'):
            response = asyncio.run(get_response(user_input))
            st.write(response)

    except Exception as e:
        st.error(f"Error processing file: {e}")